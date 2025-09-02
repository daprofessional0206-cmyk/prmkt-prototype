# shared/exports.py
from __future__ import annotations
from typing import Iterable, Optional, List
import io
import os

try:
    from fpdf import FPDF
except Exception:  # fpdf not installed
    FPDF = None

try:
    from docx import Document  # python-docx
except Exception:
    Document = None


def _find_dejavu() -> Optional[str]:
    """
    Try to locate a DejaVuSans*.ttf we saved under shared/fonts or shared/fonts/dejavu.
    Returns a file path or None.
    """
    here = os.path.dirname(__file__)
    candidates: List[str] = []
    # local bundle
    candidates.append(os.path.join(here, "fonts", "DejaVuSans.ttf"))
    # subfolder bundle (recommended)
    candidates.append(os.path.join(here, "fonts", "dejavu", "DejaVuSans.ttf"))
    # system-ish fallbacks
    candidates.append(os.path.expanduser("~/.fonts/DejaVuSans.ttf"))
    candidates.append("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")

    for p in candidates:
        if os.path.isfile(p):
            return p
    return None


def text_to_pdf_bytes(text: str, title: str = "Document") -> bytes:
    """
    Robust PDF exporter:
      - Handles Unicode (adds DejaVuSans Unicode font if available)
      - Wraps long lines
      - Avoids 'Not enough horizontal space...' error

    Requires `fpdf`. If missing, we return a tiny PDF-like message.
    """
    if FPDF is None:
        return b"PDF export requires the 'fpdf' package."

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Try to enable Unicode
    font_path = _find_dejavu()
    if font_path:
        try:
            pdf.add_font("DejaVu", "", font_path, uni=True)
            pdf.set_font("DejaVu", size=12)
        except Exception:
            pdf.set_font("Arial", size=12)
    else:
        pdf.set_font("Arial", size=12)

    try:
        pdf.set_title(title)
    except Exception:
        pass

    # Write lines with safe wrapping
    for raw in (text or "").splitlines():
        line = raw.replace("\t", "    ")
        if line.strip() == "":
            # blank line => vertical space
            pdf.ln(6)
            continue
        # width=0 means take full width minus margins; height=6 is a nice leading
        pdf.multi_cell(w=0, h=6, txt=line)

    # fpdf.output(dest="S") returns str in some versions, bytes in others
    out = pdf.output(dest="S")
    return out if isinstance(out, (bytes, bytearray)) else out.encode("latin-1", "ignore")


def text_to_docx_bytes(text: str, title: str = "Document") -> bytes:
    """
    Create a simple .docx with a heading and paragraphs.
    Requires `python-docx`. If missing, returns a .txt-ish fallback in bytes.
    """
    if Document is None:
        return (f"{title}\n\n{text}").encode("utf-8")

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for raw in (text or "").splitlines():
        if raw.strip() == "":
            doc.add_paragraph("")  # blank line
        else:
            doc.add_paragraph(raw)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def join_variants(variants: Iterable[str], divider: Optional[str] = None) -> str:
    """
    Utility: merge A/B/C variants for exporting. Adds a friendly divider.
    """
    joined: List[str] = []
    for i, v in enumerate(variants, start=1):
        title = f"Variant {i}"
        block = f"{title}\n{'-'*len(title)}\n{(v or '').strip()}"
        joined.append(block)

    if divider is None:
        divider = "\n\n" + ("-" * 60) + "\n\n"

    return divider.join(joined)
